#!/usr/bin/env python3
"""
Parse PlayENG Oxford A1 docx files into structured JSON.
Extracts vocabulary + sentences from matching table structures.
"""
import json
import re
import sys
from docx import Document

VOCAB_PATH = '../docs/OXFORD_A1_SZOSZEDET_v9u.docx'
SENTENCES_PATH = '../docs/OXFORD_A1_MONDATOK_v2.docx'
DIALOGUES_PATH = '../docs/OXFORD_A1_PARBESZEDEK_v1.docx'
OUTPUT_PATH = '../backend/data/oxford3000.json'

# 23 chapters grouped into 6 levels
CHAPTER_DEFS = [
    # Level 1: Alapok (Basics)
    {"id": 1, "name": "Alapvető kommunikáció", "nameEn": "Basic Communication", "level": 1},
    {"id": 2, "name": "Névmások és névelők", "nameEn": "Pronouns & Articles", "level": 1},
    {"id": 3, "name": "Számok és mértékek", "nameEn": "Numbers & Measurements", "level": 1},
    {"id": 4, "name": "Színek", "nameEn": "Colours", "level": 1},
    # Level 2: Mindennapok (Daily Life)
    {"id": 5, "name": "Emberek és család", "nameEn": "People & Family", "level": 2},
    {"id": 6, "name": "Test és egészség", "nameEn": "Body & Health", "level": 2},
    {"id": 7, "name": "Étel és ital", "nameEn": "Food & Drink", "level": 2},
    {"id": 8, "name": "Ruházat és megjelenés", "nameEn": "Clothing & Appearance", "level": 2},
    {"id": 9, "name": "Otthon és lakás", "nameEn": "Home & Living", "level": 2},
    # Level 3: Világ körülöttünk (Our World)
    {"id": 10, "name": "Idő", "nameEn": "Time", "level": 3},
    {"id": 11, "name": "Iskola és tanulás", "nameEn": "School & Learning", "level": 3},
    {"id": 12, "name": "Helyek és épületek", "nameEn": "Places & Buildings", "level": 3},
    {"id": 13, "name": "Közlekedés és utazás", "nameEn": "Transport & Travel", "level": 3},
    {"id": 14, "name": "Természet és időjárás", "nameEn": "Nature & Weather", "level": 3},
    # Level 4: Társalgás (Conversation)
    {"id": 15, "name": "Érzelmek és jellem", "nameEn": "Emotions & Character", "level": 4},
    {"id": 16, "name": "Szabadidő és szórakozás", "nameEn": "Leisure & Entertainment", "level": 4},
    {"id": 17, "name": "Foglalkozások", "nameEn": "Jobs", "level": 4},
    # Level 5: Kapcsolatok (Connections)
    {"id": 18, "name": "Munka és üzlet", "nameEn": "Work & Business", "level": 5},
    {"id": 19, "name": "Technológia és média", "nameEn": "Technology & Media", "level": 5},
    {"id": 20, "name": "Igék — Alapvető cselekvések", "nameEn": "Verbs — Basic Actions", "level": 5},
    # Level 6: Magabiztosság (Confidence)
    {"id": 21, "name": "Melléknevek", "nameEn": "Adjectives", "level": 6},
    {"id": 22, "name": "Határozószók, elöljárószók és kötőszók", "nameEn": "Adverbs, Prepositions & Conjunctions", "level": 6},
    {"id": 23, "name": "Gyakori szókapcsolatok", "nameEn": "Common Collocations", "level": 6},
]

LEVEL_DEFS = [
    {"id": 1, "name": "Alapok", "nameEn": "Basics", "icon": "🌱", "description": "Köszönés, névmások, számok, színek"},
    {"id": 2, "name": "Mindennapok", "nameEn": "Daily Life", "icon": "🏠", "description": "Család, test, étel, ruha, otthon"},
    {"id": 3, "name": "Világ körülöttünk", "nameEn": "Our World", "icon": "🌍", "description": "Idő, iskola, helyek, közlekedés, természet"},
    {"id": 4, "name": "Társalgás", "nameEn": "Conversation", "icon": "💬", "description": "Érzelmek, szabadidő, foglalkozások"},
    {"id": 5, "name": "Kapcsolatok", "nameEn": "Connections", "icon": "🤝", "description": "Munka, technológia, alapvető igék"},
    {"id": 6, "name": "Magabiztosság", "nameEn": "Confidence", "icon": "🎯", "description": "Melléknevek, határozók, szókapcsolatok"},
]


def extract_tables(doc_path):
    """Extract all tables from a docx, returning list of list of (col0, col1) tuples."""
    doc = Document(doc_path)
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2:
                rows.append((cells[0], cells[1]))
        tables.append(rows)
    return tables


def find_chapter_for_table(table_index, total_tables):
    """Map table indices to chapters. Tables follow the 23-chapter structure with sub-tables."""
    return table_index  # We'll map by matching content instead


def parse_vocab_and_sentences():
    """Parse vocabulary and sentences docx files into structured data."""
    print("Parsing vocabulary...")
    vocab_tables = extract_tables(VOCAB_PATH)
    print(f"  Found {len(vocab_tables)} vocabulary tables")

    print("Parsing sentences...")
    sent_tables = extract_tables(SENTENCES_PATH)
    print(f"  Found {len(sent_tables)} sentence tables")

    # Build word list from vocab tables, matching with sentence tables
    all_words = []
    word_id = 0

    # Track which chapter we're in by counting tables
    # We need to figure out table->chapter mapping
    # Let's look at the content of tables and paragraph flow

    vocab_doc = Document(VOCAB_PATH)

    # Build a mapping: for each table, find which chapter heading precedes it
    # We'll walk through the document body elements
    chapter_num = 0
    sub_section = ""
    table_chapter_map = {}
    table_idx = 0

    for element in vocab_doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'p':
            # Check for chapter heading pattern like "1. Alapvető kommunikáció"
            text = element.text or ''
            # Also check all runs
            for run in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if run.text:
                    text += run.text
            text = text.strip()
            # Match chapter heading: number. Name
            m = re.match(r'^(\d+)\.\s+', text)
            if m and int(m.group(1)) <= 23:
                chapter_num = int(m.group(1))
        elif tag == 'tbl':
            table_chapter_map[table_idx] = chapter_num
            table_idx += 1

    print(f"  Mapped {len(table_chapter_map)} tables to chapters")

    # Now build words
    for t_idx, vtable in enumerate(vocab_tables):
        chapter = table_chapter_map.get(t_idx, 0)
        if chapter == 0:
            continue

        # Get matching sentence table
        stable = sent_tables[t_idx] if t_idx < len(sent_tables) else []

        # Skip header row
        for row_idx in range(len(vtable)):
            en_word, hu_word = vtable[row_idx]
            # Skip header rows
            if en_word.lower() in ('english', 'angol', ''):
                continue
            if not en_word or not hu_word:
                continue

            word_id += 1

            # Find matching sentence
            sentences = []
            if row_idx < len(stable):
                en_sent, hu_sent = stable[row_idx]
                if en_sent and en_sent.lower() not in ('english', 'angol', ''):
                    sentences.append({"en": en_sent, "hu": hu_sent})

            # Determine POS from context (basic heuristic)
            pos = guess_pos(en_word, chapter)

            word_obj = {
                "id": word_id,
                "word": en_word.lower().strip(),
                "wordDisplay": en_word.strip(),
                "hungarian": hu_word.strip(),
                "chapter": chapter,
                "pos": pos,
                "sentences": sentences
            }
            all_words.append(word_obj)

    print(f"  Total words extracted: {len(all_words)}")
    return all_words


def guess_pos(word, chapter):
    """Guess part of speech based on chapter context."""
    if chapter == 20:
        return "verb"
    elif chapter == 21:
        return "adjective"
    elif chapter == 22:
        return "adverb/preposition"
    elif chapter == 23:
        return "collocation"
    elif chapter in (1,):
        return "exclamation"
    elif chapter in (2,):
        return "pronoun"
    elif chapter in (3,):
        return "number"
    elif chapter == 4:
        return "adjective"
    elif chapter in (5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19):
        return "noun"
    return "word"


def parse_dialogues():
    """Parse dialogues docx into structured data."""
    print("Parsing dialogues...")
    try:
        doc = Document(DIALOGUES_PATH)
        dialogues = []
        dialogue_id = 0

        current_dialogue = None
        chapter_num = 0

        for element in doc.element.body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
            if tag == 'p':
                text = ''
                for run in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                    if run.text:
                        text += run.text
                text = text.strip()

                m = re.match(r'^(\d+)\.\s+', text)
                if m and int(m.group(1)) <= 23:
                    chapter_num = int(m.group(1))
            elif tag == 'tbl':
                table_el = doc.tables[dialogue_id] if dialogue_id < len(doc.tables) else None
                if table_el:
                    rows = []
                    for row in table_el.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        rows.append(cells)

                    if len(rows) > 1:
                        dialogue_id += 1
                        dialogues.append({
                            "id": dialogue_id,
                            "chapterId": chapter_num,
                            "turns": rows
                        })

        print(f"  Found {len(dialogues)} dialogue tables")
        return dialogues
    except Exception as e:
        print(f"  Error parsing dialogues: {e}")
        return []


def build_output(words, dialogues):
    """Build the final JSON structure."""
    # Group words by chapter
    chapters = []
    for ch_def in CHAPTER_DEFS:
        ch_words = [w for w in words if w["chapter"] == ch_def["id"]]
        chapters.append({
            "id": ch_def["id"],
            "name": ch_def["name"],
            "nameEn": ch_def["nameEn"],
            "level": ch_def["level"],
            "wordCount": len(ch_words),
            "words": ch_words
        })

    output = {
        "levels": LEVEL_DEFS,
        "chapters": chapters,
        "dialogues": dialogues,
        "meta": {
            "totalWords": len(words),
            "totalChapters": len(chapters),
            "totalLevels": 6,
            "source": "Oxford 3000 A1"
        }
    }
    return output


def main():
    import os
    os.chdir(os.path.dirname(os.path.abspath(__file__)))

    words = parse_vocab_and_sentences()
    dialogues = parse_dialogues()
    output = build_output(words, dialogues)

    with open(OUTPUT_PATH, 'w', encoding='utf-8') as f:
        json.dump(output, f, ensure_ascii=False, indent=2)

    print(f"\nOutput written to {OUTPUT_PATH}")
    print(f"  Levels: {len(output['levels'])}")
    print(f"  Chapters: {len(output['chapters'])}")
    print(f"  Total words: {output['meta']['totalWords']}")
    print(f"  Dialogues: {len(output['dialogues'])}")

    # Summary per level
    for level in LEVEL_DEFS:
        level_chapters = [c for c in output['chapters'] if c['level'] == level['id']]
        level_words = sum(c['wordCount'] for c in level_chapters)
        ch_names = ', '.join(f"{c['id']}.{c['name']}" for c in level_chapters)
        print(f"  Level {level['id']} ({level['name']}): {level_words} words [{ch_names}]")


if __name__ == '__main__':
    main()
